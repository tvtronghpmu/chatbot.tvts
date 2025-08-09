import streamlit as st
import os
import openai
import pandas as pd
from docx import Document
import re
import time
from io import BytesIO
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from dotenv import load_dotenv

# --- CẤU HÌNH ---

# Tải biến môi trường từ file .env
load_dotenv()
# Khởi tạo client OpenAI (cú pháp mới)
# Đảm bảo bạn đã có file .env với nội dung: OPENAI_API_KEY="sk-..."
try:
    openai.api_key = os.getenv("OPENAI_API_KEY")
except Exception as e:
    st.error(f"Không thể khởi tạo OpenAI Client. Vui lòng kiểm tra API Key. Lỗi: {e}")
    st.stop()


# Tiêu đề và mô tả ứng dụng
st.title("🤖 Tư vấn Tuyển sinh - Đại học Y Dược Hải Phòng")
st.markdown("""
Chào mừng bạn đến với chatbot tư vấn tuyển sinh của Trường Đại học Y Dược Hải Phòng!
Hãy tải lên các tài liệu tuyển sinh (PDF, DOCX, XLSX) sau đó đặt câu hỏi về thông tin tuyển sinh.
""")

# --- CÁC HÀM XỬ LÝ FILE (ĐÃ TỐI ƯU VỚI CACHING) ---

# Gợi ý: Để code chạy được, hãy cài đặt Tesseract và thêm nó vào biến môi trường PATH của hệ thống.
# Sau đó có thể xóa dòng code gán đường dẫn cố định bên dưới.
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' # Không khuyến khích

@st.cache_data(max_entries=50) # Cache kết quả xử lý file
def read_pdf_optimized(file_bytes):
    """
    Đọc file PDF một cách tối ưu:
    1. Sử dụng PyMuPDF (fitz) để có hiệu suất cao nhất.
    2. Trích xuất văn bản và bảng.
    3. Áp dụng "Smart OCR": chỉ OCR những trang không có văn bản (trang ảnh).
    """
    text = ""
    MIN_TEXT_LENGTH_FOR_OCR = 20  # Ngưỡng ký tự để quyết định có OCR hay không

    try:
        # Mở file PDF từ bytes
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")

        for page_num, page in enumerate(pdf_doc, start=1):
            text += f"\n--- Trang {page_num} ---\n"

            # 1. Trích xuất văn bản thông thường
            page_text = page.get_text("text")
            if page_text:
                text += page_text + "\n"

            # 2. Trích xuất bảng (PyMuPDF cũng có thể làm điều này)
            tabs = page.find_tables()
            if tabs.tables:
                for i, tab in enumerate(tabs):
                    text += f"[BẢNG {i+1}]\n"
                    df = tab.to_pandas()
                    text += df.to_string(index=False) + "\n"

            # 3. Smart OCR: Chỉ thực hiện OCR nếu trang có rất ít hoặc không có văn bản
            if len(page_text.strip()) < MIN_TEXT_LENGTH_FOR_OCR:
                images = page.get_images(full=True)
                if images:
                    st.info(f"Trang {page_num} có vẻ là ảnh, đang thực hiện OCR...")
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    try:
                        # Sử dụng Tesseract để OCR với ngôn ngữ tiếng Việt
                        ocr_text = pytesseract.image_to_string(image, lang='vie')
                        if ocr_text.strip():
                           text += f"\n[OCR - Trang {page_num} - Hình {img_index+1}]\n{ocr_text}\n"
                    except Exception as ocr_error:
                        st.warning(f"Lỗi khi OCR hình ảnh trên trang {page_num}: {ocr_error}")

    except Exception as e:
        return f"Lỗi khi đọc file PDF: {e}"
        
    return text

@st.cache_data(max_entries=50)
def read_docx(file_bytes):
    """Đọc nội dung từ file DOCX."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        
        for i, table in enumerate(doc.tables):
            text += f"\n\n[BẢNG {i + 1}]\n"
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(" | ".join(row_data))
            text += "\n".join(table_data)
        return text
    except Exception as e:
        return f"Lỗi khi đọc file DOCX: {e}"


@st.cache_data(max_entries=50)
def read_xlsx(file_bytes):
    """Đọc nội dung từ file XLSX."""
    try:
        df = pd.read_excel(io.BytesIO(file_bytes))
        return df.to_string(index=False)
    except Exception as e:
        return f"Lỗi khi đọc file XLSX: {e}"

def process_uploaded_files(uploaded_files):
    """Xử lý danh sách các file được tải lên."""
    all_text = ""
    for file in uploaded_files:
        # Đọc bytes của file một lần duy nhất
        file_bytes = file.getvalue()
        file_type = file.type
        
        text = ""
        if file_type == "application/pdf":
            text = read_pdf_optimized(file_bytes)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = read_docx(file_bytes)
        elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            text = read_xlsx(file_bytes)
        else:
            text = f"Định dạng file không được hỗ trợ: {file.name}"
        
        all_text += f"\n\n--- NỘI DUNG TỪ {file.name} ---\n\n{text}"
    
    return all_text

# --- HÀM GỌI API ---
def ask_openai(question, context):
    """Gửi câu hỏi đến OpenAI và nhận câu trả lời (sử dụng cú pháp mới)."""
    prompt = f"""
    Bạn là một chatbot tư vấn tuyển sinh chuyên nghiệp và thân thiện của Trường Đại học Y Dược Hải Phòng.
    Nhiệm vụ của bạn là trả lời các câu hỏi của sinh viên và phụ huynh dựa trên nội dung các tài liệu tuyển sinh được cung cấp dưới đây.
    - Trả lời một cách chính xác, rõ ràng và đi thẳng vào vấn đề.
    - Nếu thông tin không có trong tài liệu, hãy trả lời một cách lịch sự: 'Tôi không tìm thấy thông tin bạn hỏi trong các tài liệu đã được cung cấp. Bạn vui lòng liên hệ phòng tuyển sinh để được hỗ trợ chính xác nhất.'
    - KHÔNG tự bịa đặt thông tin.
    - Thông tin so sánh liệt kê dưới dạng bảng.
    Dưới đây là toàn bộ thông tin tuyển sinh bạn có:
    ---
    {context}
    ---
    
    Câu hỏi của người dùng: {question}
    
    Câu trả lời của bạn:
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Bạn là một trợ lý ảo tư vấn tuyển sinh của Đại học Y Dược Hải Phòng."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1024,  # Tăng giới hạn token để có câu trả lời đầy đủ hơn
            temperature=0.7   # Giảm temperature để câu trả lời bám sát dữ liệu hơn
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Lỗi khi kết nối với OpenAI: {str(e)}"

# --- GIAO DIỆN STREAMLIT ---

# Sidebar
with st.sidebar:
    st.header("📂 Tải lên tài liệu")
    uploaded_files = st.file_uploader(
        "Chọn file PDF, DOCX, XLSX", 
        type=["pdf", "docx", "xlsx"],
        accept_multiple_files=True
    )
    
    st.markdown("---")
    st.markdown("**Nhóm thực hiện:** Trần Văn Trọng, Nguyễn Thanh Hà")
   # st.info("Chatbot sử dụng OpenAI API để phân tích dữ liệu và trả lời câu hỏi.")

# Xử lý và lưu trữ context
if "data_context" not in st.session_state:
    st.session_state.data_context = ""

if uploaded_files:
    # Chỉ xử lý lại nếu file tải lên thay đổi
    # So sánh tên file và kích thước để kiểm tra sự thay đổi
    uploaded_file_signatures = [(f.name, f.size) for f in uploaded_files]
    if st.session_state.get("last_uploaded_files") != uploaded_file_signatures:
        with st.spinner(f"Đang xử lý {len(uploaded_files)} file... Thao tác này có thể mất một lúc với các file lớn hoặc file scan."):
            data_context = process_uploaded_files(uploaded_files)
            st.session_state.data_context = data_context
            st.session_state.last_uploaded_files = uploaded_file_signatures
            st.success(f"Đã xử lý xong {len(uploaded_files)} file! Dữ liệu đã sẵn sàng.")
            # Xóa lịch sử chat cũ khi có dữ liệu mới
            st.session_state.messages = [
                {"role": "assistant", "content": "Dữ liệu mới đã được tải lên! Tôi có thể giúp gì cho bạn về thông tin tuyển sinh?"}
            ]
            st.rerun() # Tải lại trang để hiển thị chat mới

# Tùy chọn xem nội dung đã xử lý
if st.session_state.data_context:
    with st.expander("Xem nội dung dữ liệu đã được xử lý"):
        st.text_area("Nội dung:", st.session_state.data_context, height=300)

# Giao diện chat
st.header("💬 Hỏi đáp về tuyển sinh")

# Khởi tạo lịch sử chat
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "content": "Xin chào! Vui lòng tải tài liệu tuyển sinh lên ở thanh bên trái để bắt đầu."}
    ]

# Hiển thị lịch sử chat
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Xử lý câu hỏi mới
if user_input := st.chat_input("Nhập câu hỏi của bạn ở đây..."):
    # Thêm câu hỏi vào lịch sử
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)
    
    # Kiểm tra xem đã có dữ liệu chưa
    if not st.session_state.data_context:
        with st.chat_message("assistant"):
            st.error("⚠️ Vui lòng tải lên tài liệu tuyển sinh trước khi đặt câu hỏi!")
        st.session_state.messages.append({"role": "assistant", "content": "⚠️ Vui lòng tải lên tài liệu tuyển sinh trước khi đặt câu hỏi!"})
    else:
        # Tạo câu trả lời
        with st.chat_message("assistant"):
            with st.spinner("🧠 Đang suy nghĩ..."):
                assistant_response = ask_openai(user_input, st.session_state.data_context)
                
                # Hiệu ứng gõ máy cho đẹp mắt
                message_placeholder = st.empty()
                full_response = ""
                for chunk in assistant_response.split():
                    full_response += chunk + " "
                    time.sleep(0.02)
                    message_placeholder.markdown(full_response + "▌")
                message_placeholder.markdown(full_response)
        
        # Thêm câu trả lời vào lịch sử

        st.session_state.messages.append({"role": "assistant", "content": full_response})

